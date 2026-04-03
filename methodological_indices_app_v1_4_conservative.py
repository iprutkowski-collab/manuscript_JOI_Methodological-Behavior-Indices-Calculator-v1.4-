
from __future__ import annotations

import json
import re
import sys
import os
import tempfile
from dataclasses import dataclass, asdict
from pathlib import Path
from typing import Dict, List, Tuple, Optional

import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from tkinter.scrolledtext import ScrolledText

try:
    from docx import Document  # type: ignore
except Exception:
    Document = None

try:
    import fitz  # type: ignore  # PyMuPDF
except Exception:
    fitz = None

try:
    from pypdf import PdfReader  # type: ignore
except Exception:
    PdfReader = None

try:
    import pdfplumber  # type: ignore
except Exception:
    pdfplumber = None

try:
    from PIL import Image  # type: ignore
except Exception:
    Image = None

try:
    import pytesseract  # type: ignore
except Exception:
    pytesseract = None

try:
    from striprtf.striprtf import rtf_to_text  # type: ignore
except Exception:
    rtf_to_text = None

APP_TITLE = "Methodological Behavior Indices Calculator"
APP_VERSION = "1.4"

SIGNAL_PATTERNS: Dict[str, List[str]] = {
    "sig_likert": [r"\blikert\b", r"\blikert-type\b", r"\blikert type\b"],
    "sig_scale_points": [r"\b5-point\b", r"\b7-point\b", r"\b10-point\b", r"\bpoint scale\b"],
    "sig_scale_instrument": [r"\bscale\b", r"\binstrument\b", r"\bmeasurement scale\b", r"\bvisual analogue scale\b", r"\bvas\b", r"\bnrs\b"],
    "sig_strongly_agree": [r"\bstrongly agree\b", r"\bstrongly disagree\b"],
    "sig_ordinal": [r"\bordinal\b", r"\binterval\b", r"\bratio\b"],
    "sig_mean_sd": [r"\bmean\b", r"\bmeans\b", r"\bstandard deviation\b", r"\bsd\b", r"\bm\s*=\s*\d"],
    "sig_median_iqr": [r"\bmedian\b", r"\binterquartile range\b", r"\biqr\b"],
    "sig_ttest_anova": [r"\bt-test\b", r"\bttest\b", r"\banova\b", r"\bmanova\b", r"\bmann[–-]?whitney\b", r"\bkruskal[–-]?wallis\b"],
    "sig_regression": [r"\bregression\b", r"\blogistic regression\b", r"\blinear regression\b"],
    "sig_sem_cfa": [r"\bsem\b", r"\bcfa\b", r"\bstructural equation model(?:ling)?\b", r"\bconfirmatory factor analysis\b", r"\bchi-square\b"],
    "sig_nonparametric": [r"\bnonparametric\b", r"\bparametric assumption\b", r"\bassumption\b"],
    "sig_cronbach_alpha": [r"\bcronbach'?s alpha\b", r"\balpha\s*=\s*\d", r"\bα\s*=\s*\d"],
    "sig_reliability_validity": [r"\breliability\b", r"\bvalidity\b", r"\bave\b", r"\bcomposite reliability\b", r"\bconvergent validity\b", r"\bdiscriminant validity\b", r"\bfactor loading\b"],
    "sig_survey_questionnaire": [r"\bsurvey\b", r"\bquestionnaire\b", r"\bpilot test\b", r"\bpretest\b", r"\bvalidation study\b", r"\bexploratory factor analysis\b", r"\brobustness check\b"],
}

CLASS_MAP: Dict[str, List[str]] = {
    "Scale / instrument signals": ["sig_likert", "sig_scale_points", "sig_scale_instrument", "sig_strongly_agree"],
    "Measurement-level signals": ["sig_ordinal"],
    "Descriptive-statistics signals": ["sig_mean_sd", "sig_median_iqr"],
    "Inferential-technique signals": ["sig_ttest_anova", "sig_regression", "sig_sem_cfa"],
    "Analytical-treatment / assumption signals": ["sig_nonparametric"],
    "Reliability and validity signals": ["sig_cronbach_alpha", "sig_reliability_validity"],
    "Other methodological signals": ["sig_survey_questionnaire"],
}

AMBIGUOUS_PATTERNS = [r"\bscale\b", r"\binstrument\b", r"\bsurvey\b", r"\bquestionnaire\b"]
DEFAULT_WEIGHTS = (0.4, 0.4, 0.2)

@dataclass
class ComputationResult:
    file_path: str
    references_excluded: bool
    word_count_excluding_references: int
    total_hit_count: int
    ambiguous_hit_count: int
    active_signal_count: int
    total_operational_signals: int
    active_class_count: int
    total_signal_classes: int
    signal_flags: Dict[str, int]
    class_flags: Dict[str, int]
    signal_hits: Dict[str, int]
    bundle_membership: Dict[str, float]
    bundle_label: str
    spi: float
    mvi: float
    bcs: float
    aus: float
    loader_used: str
    loader_note: str
    text_excerpt_note: str

def normalize_text(text: str) -> str:
    text = text.replace("\u2013", "-").replace("\u2014", "-")
    text = re.sub(r"\s+", " ", text)
    return text.strip()

def remove_references_section(text: str) -> Tuple[str, bool]:
    patterns = [
        r"\n\s*references\s*\n",
        r"\n\s*bibliography\s*\n",
        r"\n\s*works cited\s*\n",
    ]
    lower = "\n" + text.lower() + "\n"
    positions = []
    for pat in patterns:
        match = re.search(pat, lower, flags=re.IGNORECASE)
        if match:
            positions.append(match.start())
    if positions:
        cut = min(positions)
        return text[: max(cut - 1, 0)].strip(), True
    return text, False

def _read_plain_text(path: str) -> str:
    for enc in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return Path(path).read_text(encoding=enc, errors="ignore")
        except Exception:
            continue
    return Path(path).read_text(errors="ignore")

def _read_docx(path: str) -> str:
    if Document is None:
        raise RuntimeError("python-docx is not installed. DOCX reading is unavailable.")
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.append("\t".join(cell.text for cell in row.cells))
    return "\n".join(parts)

def _read_doc_via_word(path: str) -> str:
    if not sys.platform.startswith("win"):
        raise RuntimeError("Legacy .doc reading is only supported on Windows.")
    try:
        import win32com.client  # type: ignore
    except Exception:
        raise RuntimeError("Legacy .doc support requires pywin32 (`pip install pywin32`) and Microsoft Word installed.")
    word = None
    temp_docx = None
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(path, ReadOnly=True)
        fd, tmpname = tempfile.mkstemp(suffix=".docx")
        os.close(fd)
        temp_docx = tmpname
        doc.SaveAs(temp_docx, FileFormat=16)  # wdFormatXMLDocument
        doc.Close(False)
        return _read_docx(temp_docx)
    finally:
        if word is not None:
            try:
                word.Quit()
            except Exception:
                pass
        if temp_docx and os.path.exists(temp_docx):
            try:
                os.remove(temp_docx)
            except Exception:
                pass

def _read_rtf(path: str) -> str:
    raw = _read_plain_text(path)
    if rtf_to_text is not None:
        try:
            return rtf_to_text(raw)
        except Exception:
            pass
    # very simple fallback
    text = re.sub(r"{\\.*?}", " ", raw)
    text = re.sub(r"\\'[0-9a-fA-F]{2}", " ", text)
    text = re.sub(r"\\[a-zA-Z]+\d* ?", " ", text)
    text = text.replace("{", " ").replace("}", " ")
    return text

def _ocr_pdf_via_pymupdf(path: str) -> str:
    if fitz is None or pytesseract is None or Image is None:
        return ""
    try:
        doc = fitz.open(path)
        texts = []
        for page in doc:
            pix = page.get_pixmap(dpi=200, alpha=False)
            mode = "RGB"
            img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)
            txt = pytesseract.image_to_string(img)
            if txt.strip():
                texts.append(txt)
        doc.close()
        return "\n".join(texts)
    except Exception:
        return ""

def load_text_from_file(path: str) -> Tuple[str, str, str]:
    ext = Path(path).suffix.lower()
    if ext in {".txt", ".md", ".csv"}:
        return _read_plain_text(path), "plain_text", "Loaded as plain text."
    if ext == ".docx":
        return _read_docx(path), "docx", "Loaded with python-docx."
    if ext == ".doc":
        return _read_doc_via_word(path), "doc_via_word", "Loaded from legacy .doc via Microsoft Word automation."
    if ext == ".rtf":
        return _read_rtf(path), "rtf", "Loaded from RTF."
    if ext == ".pdf":
        errors = []
        if fitz is not None:
            try:
                doc = fitz.open(path)
                txt = "\n".join(page.get_text("text") or "" for page in doc)
                doc.close()
                if txt.strip():
                    return txt, "pdf_pymupdf", "Loaded PDF text with PyMuPDF."
            except Exception as e:
                errors.append(f"PyMuPDF: {e}")
        if PdfReader is not None:
            try:
                reader = PdfReader(path)
                txt = "\n".join((page.extract_text() or "") for page in reader.pages)
                if txt.strip():
                    return txt, "pdf_pypdf", "Loaded PDF text with pypdf."
            except Exception as e:
                errors.append(f"pypdf: {e}")
        if pdfplumber is not None:
            try:
                with pdfplumber.open(path) as pdf:
                    txt = "\n".join((page.extract_text() or "") for page in pdf.pages)
                if txt.strip():
                    return txt, "pdf_pdfplumber", "Loaded PDF text with pdfplumber."
            except Exception as e:
                errors.append(f"pdfplumber: {e}")
        ocr_text = _ocr_pdf_via_pymupdf(path)
        if ocr_text.strip():
            return ocr_text, "pdf_ocr", "Loaded PDF via OCR fallback."
        if fitz is None and PdfReader is None and pdfplumber is None:
            raise RuntimeError("PDF support is unavailable. Install PyMuPDF, pypdf, or pdfplumber.")
        msg = "The PDF could be opened, but no readable text was extracted."
        msg += " The file may be a scanned image PDF without an embedded text layer."
        msg += " Use OCR or provide a DOCX/TXT version."
        if pytesseract is None:
            msg += " Optional OCR support can be added by installing pytesseract and Tesseract OCR."
        if errors:
            msg += " Details: " + "; ".join(errors)
        raise RuntimeError(msg)
    raise RuntimeError(f"Unsupported file type: {ext}. Supported: PDF, DOCX, DOC, TXT, MD, CSV, RTF.")

def count_words(text: str) -> int:
    return len(re.findall(r"\b\w+\b", text))

def count_hits(text: str, patterns: List[str]) -> int:
    hits = 0
    for pat in patterns:
        hits += len(re.findall(pat, text, flags=re.IGNORECASE))
    return hits

def compute_indices(file_path: str, exclude_references: bool = True, weights: Tuple[float, float, float] = DEFAULT_WEIGHTS) -> ComputationResult:
    raw, loader_used, loader_note = load_text_from_file(file_path)
    raw = raw.replace("\r\n", "\n").replace("\r", "\n")
    processed = raw
    references_removed = False
    if exclude_references:
        processed, references_removed = remove_references_section(processed)
    processed = normalize_text(processed)

    word_count = count_words(processed)
    signal_flags: Dict[str, int] = {}
    signal_hits: Dict[str, int] = {}

    for signal, patterns in SIGNAL_PATTERNS.items():
        hits = count_hits(processed, patterns)
        signal_hits[signal] = hits
        signal_flags[signal] = 1 if hits > 0 else 0

    class_flags: Dict[str, int] = {}
    for cls, signals in CLASS_MAP.items():
        class_flags[cls] = 1 if any(signal_flags[s] for s in signals) else 0

    n_active = sum(signal_flags.values())
    c_active = sum(class_flags.values())
    total_hits = sum(signal_hits.values())
    ambiguous_hits = count_hits(processed, AMBIGUOUS_PATTERNS)

    likert_side = (
        signal_flags["sig_likert"]
        + signal_flags["sig_scale_points"]
        + signal_flags["sig_scale_instrument"]
        + signal_flags["sig_strongly_agree"]
    ) / 4.0
    parametric_side = (
        signal_flags["sig_mean_sd"]
        + signal_flags["sig_ttest_anova"]
        + signal_flags["sig_regression"]
        + signal_flags["sig_sem_cfa"]
    ) / 4.0
    # In version 1.4, sig_ordinal is treated conservatively.
    # If ordinal measurement-level language is visible, the record is not assigned
    # to a pure parametric profile. This avoids implying methodological correctness.
    nonparametric_core_side = (
        signal_flags["sig_median_iqr"]
        + signal_flags["sig_nonparametric"]
    ) / 2.0
    nonparametric_visibility_side = (
        signal_flags["sig_ordinal"]
        + signal_flags["sig_median_iqr"]
        + signal_flags["sig_nonparametric"]
    ) / 3.0

    ordinal_present = signal_flags["sig_ordinal"] == 1
    param_unique = (
        likert_side > 0
        and parametric_side > 0
        and nonparametric_core_side == 0
        and not ordinal_present
    )
    nonparam_unique = likert_side > 0 and nonparametric_core_side > 0 and parametric_side == 0
    other_mixed = likert_side > 0 and not (param_unique or nonparam_unique)

    membership_param = (likert_side + parametric_side) / 2.0 if param_unique else 0.0
    membership_nonparam = (likert_side + nonparametric_visibility_side) / 2.0 if nonparam_unique else 0.0
    membership_other_mixed = 0.0
    if other_mixed:
        if ordinal_present and parametric_side > 0 and nonparametric_core_side == 0:
            membership_other_mixed = (likert_side + parametric_side) / 2.0
        elif parametric_side > 0 and nonparametric_core_side > 0:
            membership_other_mixed = max((likert_side + parametric_side) / 2.0, (likert_side + nonparametric_visibility_side) / 2.0)
        elif parametric_side > 0:
            membership_other_mixed = (likert_side + parametric_side) / 2.0
        elif nonparametric_core_side > 0:
            membership_other_mixed = (likert_side + nonparametric_visibility_side) / 2.0
        else:
            membership_other_mixed = likert_side

    bundle_membership = {
        "bundle_likert_parametric_profile": round(membership_param, 4),
        "bundle_likert_nonparametric_profile": round(membership_nonparam, 4),
        "bundle_likert_other_mixed": round(membership_other_mixed, 4),
    }
    bundle_label = max(bundle_membership.items(), key=lambda kv: kv[1])[0] if max(bundle_membership.values()) > 0 else "no_likert_bundle_detected"

    w1, w2, w3 = weights
    K = len(SIGNAL_PATTERNS)
    C = len(CLASS_MAP)
    spi = n_active / K if K else 0.0
    density = total_hits / (word_count / 1000.0) if word_count > 0 else 0.0
    mvi = w1 * (n_active / K if K else 0.0) + w2 * (c_active / C if C else 0.0) + w3 * density
    bcs = max(bundle_membership.values()) if bundle_membership else 0.0
    aus = ambiguous_hits / total_hits if total_hits > 0 else 0.0

    return ComputationResult(
        file_path=file_path,
        references_excluded=references_removed if exclude_references else False,
        word_count_excluding_references=word_count,
        total_hit_count=total_hits,
        ambiguous_hit_count=ambiguous_hits,
        active_signal_count=n_active,
        total_operational_signals=K,
        active_class_count=c_active,
        total_signal_classes=C,
        signal_flags=signal_flags,
        class_flags=class_flags,
        signal_hits=signal_hits,
        bundle_membership=bundle_membership,
        bundle_label=bundle_label,
        spi=round(spi, 4),
        mvi=round(mvi, 4),
        bcs=round(bcs, 4),
        aus=round(aus, 4),
        loader_used=loader_used,
        loader_note=loader_note,
        text_excerpt_note="Word count excludes References when a References/Bibliography heading is detected.",
    )

def build_report_text(result: ComputationResult) -> str:
    lines: List[str] = []
    lines.append(f"{APP_TITLE} — Report")
    lines.append("")
    lines.append(f"File: {result.file_path}")
    lines.append(f"Loader: {result.loader_used}")
    lines.append(f"Loader note: {result.loader_note}")
    lines.append(f"References excluded: {'Yes' if result.references_excluded else 'No / not detected'}")
    lines.append(result.text_excerpt_note)
    lines.append("")
    lines.append("Core values")
    lines.append(f"- Word count: {result.word_count_excluding_references}")
    lines.append(f"- Total methodological hits: {result.total_hit_count}")
    lines.append(f"- Ambiguous hits: {result.ambiguous_hit_count}")
    lines.append(f"- Active signals: {result.active_signal_count}/{result.total_operational_signals}")
    lines.append(f"- Active classes: {result.active_class_count}/{result.total_signal_classes}")
    lines.append("")
    lines.append("Indices")
    lines.append(f"- SPI = {result.spi:.4f}")
    lines.append(f"- MVI = {result.mvi:.4f}")
    lines.append(f"- BCS = {result.bcs:.4f}")
    lines.append(f"- AUS = {result.aus:.4f}")
    lines.append(f"- Dominant bundle = {result.bundle_label}")
    lines.append("")
    lines.append("Signal flags")
    for key, val in result.signal_flags.items():
        lines.append(f"- {key}: {val} (hits={result.signal_hits[key]})")
    lines.append("")
    lines.append("Class flags")
    for key, val in result.class_flags.items():
        lines.append(f"- {key}: {val}")
    lines.append("")
    lines.append("Bundle membership")
    for key, val in result.bundle_membership.items():
        lines.append(f"- {key}: {val:.4f}")
    return "\n".join(lines)

class App(tk.Tk):
    def __init__(self) -> None:
        super().__init__()
        self.title(f"{APP_TITLE} v{APP_VERSION}")
        self.geometry("1200x780")
        self.minsize(1020, 700)
        self.result: Optional[ComputationResult] = None
        self._build_ui()

    def _build_ui(self) -> None:
        top = ttk.Frame(self, padding=10)
        top.pack(fill=tk.X)

        self.file_var = tk.StringVar()
        ttk.Label(top, text="Article file:").grid(row=0, column=0, sticky="w")
        ttk.Entry(top, textvariable=self.file_var, width=90).grid(row=0, column=1, padx=6, sticky="ew")
        ttk.Button(top, text="Browse", command=self._browse).grid(row=0, column=2, padx=4)
        ttk.Button(top, text="Calculate", command=self._calculate).grid(row=0, column=3, padx=4)
        ttk.Button(top, text="Export report", command=self._export_report).grid(row=0, column=4, padx=4)
        top.columnconfigure(1, weight=1)

        self.exclude_refs_var = tk.BooleanVar(value=True)
        ttk.Checkbutton(top, text="Exclude References/Bibliography if detected", variable=self.exclude_refs_var).grid(row=1, column=1, sticky="w", pady=(8, 0))

        note = ttk.Label(
            top,
            text="Supports: PDF, DOCX, DOC (optional on Windows with Word), TXT, MD, CSV, RTF. Scanned PDFs require OCR.",
            foreground="#444444",
        )
        note.grid(row=2, column=1, sticky="w", pady=(6, 0))

        paned = ttk.Panedwindow(self, orient=tk.HORIZONTAL)
        paned.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)

        left = ttk.Frame(paned)
        right = ttk.Frame(paned)
        paned.add(left, weight=3)
        paned.add(right, weight=2)

        nb = ttk.Notebook(left)
        nb.pack(fill=tk.BOTH, expand=True)

        self.summary = ScrolledText(nb, wrap=tk.WORD, font=("Consolas", 10))
        self.interpretation = ScrolledText(nb, wrap=tk.WORD, font=("Segoe UI", 10))
        nb.add(self.summary, text="Report")
        nb.add(self.interpretation, text="Interpretation")

        ttk.Label(right, text="Detected signals and classes", font=("Segoe UI", 11, "bold")).pack(anchor="w")
        self.tree = ttk.Treeview(right, columns=("flag", "hits"), show="tree headings", height=22)
        self.tree.heading("#0", text="Name")
        self.tree.heading("flag", text="Active")
        self.tree.heading("hits", text="Hits")
        self.tree.column("#0", width=380)
        self.tree.column("flag", width=80, anchor="center")
        self.tree.column("hits", width=80, anchor="center")
        self.tree.pack(fill=tk.BOTH, expand=True, pady=(6, 0))

        status_frame = ttk.Frame(self, padding=(10, 0, 10, 10))
        status_frame.pack(fill=tk.X)
        self.status_var = tk.StringVar(value="Ready.")
        ttk.Label(status_frame, textvariable=self.status_var).pack(anchor="w")

    def _browse(self) -> None:
        path = filedialog.askopenfilename(
            title="Select article",
            filetypes=[
                ("Supported files", "*.pdf *.docx *.doc *.txt *.md *.csv *.rtf"),
                ("PDF file", "*.pdf"),
                ("Word document", "*.docx;*.doc"),
                ("Text file", "*.txt;*.md;*.csv"),
                ("RTF file", "*.rtf"),
                ("All files", "*.*"),
            ],
        )
        if path:
            self.file_var.set(path)

    def _calculate(self) -> None:
        path = self.file_var.get().strip()
        if not path:
            messagebox.showwarning(APP_TITLE, "Please choose an article file first.")
            return
        try:
            self.result = compute_indices(path, exclude_references=self.exclude_refs_var.get())
        except Exception as exc:
            messagebox.showerror(APP_TITLE, str(exc))
            self.status_var.set("Calculation failed.")
            return

        self.summary.delete("1.0", tk.END)
        self.summary.insert(tk.END, build_report_text(self.result))

        self.interpretation.delete("1.0", tk.END)
        self.interpretation.insert(tk.END, self._build_interpretation(self.result))

        for item in self.tree.get_children():
            self.tree.delete(item)
        sig_root = self.tree.insert("", tk.END, text="Operational signals")
        for name in SIGNAL_PATTERNS:
            self.tree.insert(sig_root, tk.END, text=name, values=(self.result.signal_flags[name], self.result.signal_hits[name]))
        class_root = self.tree.insert("", tk.END, text="Signal classes")
        for name, flag in self.result.class_flags.items():
            hits = sum(self.result.signal_hits[s] for s in CLASS_MAP[name])
            self.tree.insert(class_root, tk.END, text=name, values=(flag, hits))
        bundle_root = self.tree.insert("", tk.END, text="Bundle membership")
        for name, val in self.result.bundle_membership.items():
            self.tree.insert(bundle_root, tk.END, text=name, values=(f"{val:.4f}", ""))
        self.tree.item(sig_root, open=True)
        self.tree.item(class_root, open=True)
        self.tree.item(bundle_root, open=True)
        self.status_var.set(f"Calculation completed using {self.result.loader_used}.")

    def _build_interpretation(self, result: ComputationResult) -> str:
        parts = []
        parts.append("Interpretation summary")
        parts.append("")
        parts.append(f"Loader used: {result.loader_used}")
        parts.append(result.loader_note)
        parts.append("")
        parts.append(f"SPI = {result.spi:.4f}: breadth of visible methodological traces.")
        parts.append(f"MVI = {result.mvi:.4f}: overall intensity of methodological visibility.")
        parts.append(f"BCS = {result.bcs:.4f}: coherence of the detected bundle profile.")
        parts.append(f"AUS = {result.aus:.4f}: uncertainty qualifier for the visible profile.")
        parts.append("")
        parts.append(f"Dominant bundle: {result.bundle_label}.")
        parts.append("Word count excludes References when a recognizable References/Bibliography heading is detected.")
        parts.append("These results are visibility-oriented and should not be read as automatic evidence of methodological correctness.")
        return "\n".join(parts)

    def _export_report(self) -> None:
        if self.result is None:
            messagebox.showwarning(APP_TITLE, "Please calculate the indices first.")
            return
        out = filedialog.asksaveasfilename(
            title="Save report",
            defaultextension=".txt",
            filetypes=[("Text report", "*.txt"), ("JSON", "*.json")],
        )
        if not out:
            return
        try:
            if out.lower().endswith(".json"):
                with open(out, "w", encoding="utf-8") as f:
                    json.dump(asdict(self.result), f, indent=2)
            else:
                with open(out, "w", encoding="utf-8") as f:
                    f.write(build_report_text(self.result))
                    f.write("\n\n")
                    f.write(self._build_interpretation(self.result))
        except Exception as exc:
            messagebox.showerror(APP_TITLE, f"Export failed: {exc}")
            return
        self.status_var.set(f"Report saved to {out}")

def main() -> None:
    app = App()
    app.mainloop()

if __name__ == "__main__":
    main()
